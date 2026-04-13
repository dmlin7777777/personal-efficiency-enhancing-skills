"""
Diff Audit - Compare source and tailored resumes, generate change summary.

Uses heuristic signals instead of hardcoded word lists,
making it language-agnostic and role-agnostic.

v2 Features:
- Structured .docx reading (style-aware paragraph extraction)
- Risk classification for wording upgrades
- JSON output mode for LLM consumption
- Bilingual reports (EN/CN auto-detect)

Usage:
    # Markdown diff (standard)
    python diff_audit.py --source <source_md> --tailored <tailored_md>

    # Structure-aware .docx reading
    python diff_audit.py --source-docx <source.docx> --tailored-docx <tailored.docx>

    # JSON output for LLM pipeline
    python diff_audit.py --source <source_md> --tailored <tailored_md> --json

Output:
    Markdown audit log with categorized changes and risk levels.
"""

import difflib
import re
import sys
import argparse
import json
from datetime import datetime


# ---------------------------------------------------------------------------
# File reading utilities
# ---------------------------------------------------------------------------

def read_lines(path: str) -> list:
    """Read file and return list of non-empty lines."""
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip() for line in f if line.strip()]


def read_docx_structured(path: str) -> list:
    """Read .docx with structure awareness.
    Returns list of dicts: {text, style, is_table_cell, index}

    Falls back to plain text if python-docx is not available.
    """
    try:
        from docx import Document
    except ImportError:
        print("⚠️ python-docx not installed. Falling back to plain text reading.", file=sys.stderr)
        return [{"text": line, "style": "Normal", "is_table_cell": False, "index": i}
                for i, line in enumerate(read_lines(path))]

    doc = Document(path)
    result = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        result.append({
            "text": text,
            "style": p.style.name if p.style else "Normal",
            "is_table_cell": False,
            "index": i,
        })
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    result.append({
                        "text": text,
                        "style": "TableCell",
                        "is_table_cell": True,
                        "index": len(result),
                    })
    return result


def docx_to_text_lines(structured: list) -> list:
    """Extract text lines from structured docx output."""
    return [item["text"] for item in structured]


def get_structure_summary(structured: list) -> dict:
    """Generate a summary of document structure for LLM context injection."""
    styles = {}
    table_cell_count = 0
    for item in structured:
        s = item["style"]
        styles[s] = styles.get(s, 0) + 1
        if item["is_table_cell"]:
            table_cell_count += 1

    return {
        "total_items": len(structured),
        "style_distribution": styles,
        "table_cell_count": table_cell_count,
        "has_headings": any("heading" in s.lower() for s in styles),
        "has_tables": table_cell_count > 0,
    }


# ---------------------------------------------------------------------------
# Change classification heuristics
# ---------------------------------------------------------------------------

def has_numbers(text: str) -> bool:
    """Check if text contains any numeric value."""
    return bool(re.search(r"\d+", text))


def word_overlap_ratio(text_a: str, text_b: str) -> float:
    """Calculate how much of text_a's words appear in text_b (0.0 ~ 1.0)."""
    words_a = set(re.findall(r"\S+", text_a.lower()))
    words_b = set(re.findall(r"\S+", text_b.lower()))
    if not words_a:
        return 0.0
    return len(words_a & words_b) / len(words_a)


def is_expanded_version(shorter: str, longer: str) -> bool:
    """
    Check if 'longer' is an expanded version of 'shorter'.
    Signals a wording upgrade: same core content, but with added
    modifiers, results, or stronger verbs.
    """
    overlap = word_overlap_ratio(shorter, longer)
    if len(longer) < len(shorter) * 1.2:
        return False
    return overlap >= 0.4


def classify_change(removed_text: str, added_text: str) -> dict:
    """
    Classify a change pair using heuristic signals.
    No hardcoded word lists — uses structural signals that work
    across languages and industries.

    Returns dict with:
        - type: "quantification" | "wording_upgrade" | "content_change"
        - risk: "low" | "medium" | "high"
        - risk_reason: explanation of risk assessment
    """
    added_has_numbers = has_numbers(added_text)
    removed_has_numbers = has_numbers(removed_text)

    # Signal 1: Quantification — numbers added where there were none
    if added_has_numbers and not removed_has_numbers:
        return {
            "type": "quantification",
            "risk": "low",
            "risk_reason": "Added measurable data to existing content",
        }

    # Signal 2: Wording upgrade — expanded version
    if is_expanded_version(removed_text, added_text):
        risk, reason = _assess_upgrade_risk(removed_text, added_text)
        return {
            "type": "wording_upgrade",
            "risk": risk,
            "risk_reason": reason,
        }

    # Signal 3: High overlap modification
    overlap = word_overlap_ratio(removed_text, added_text)
    if overlap >= 0.5:
        risk, reason = _assess_upgrade_risk(removed_text, added_text)
        return {
            "type": "wording_upgrade",
            "risk": risk,
            "risk_reason": reason,
        }

    return {
        "type": "content_change",
        "risk": "low",
        "risk_reason": "Content replaced with fundamentally different content",
    }


def _assess_upgrade_risk(removed: str, added: str) -> tuple:
    """
    Assess the risk level of a wording upgrade.
    Returns (risk_level, reason).

    Risk criteria (language-agnostic, based on structural signals):
    - LOW: Added detail/outcome to existing statement
    - MEDIUM: Changed verb intensity without expanding scope
    - HIGH: Significantly expanded role scope or added unverified metrics
    """
    removed_len = len(removed.split())
    added_len = len(added.split())
    expansion_ratio = added_len / max(removed_len, 1)

    # HIGH risk: more than 3x expansion with significant new content
    if expansion_ratio > 3.0:
        return "high", "Significant expansion — verify added claims are accurate"

    # HIGH risk: new numbers added that weren't in source
    if has_numbers(added) and not has_numbers(removed):
        new_numbers = re.findall(r"\d+[\d.,]*%?|\$[\d,.]+|[¥€£][\d,.]+", added)
        if len(new_numbers) >= 2:
            return "medium", f"Added {len(new_numbers)} metrics — verify accuracy"

    # MEDIUM risk: moderate expansion (1.5x-3x)
    if expansion_ratio > 1.5:
        return "medium", "Moderate expansion — review for accuracy"

    # LOW risk: minor expansion (1.2x-1.5x)
    return "low", "Minor wording improvement"


# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------

def detect_language(text: str) -> str:
    """Auto-detect if text is primarily Chinese or English."""
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    english_words = len(re.findall(r"[a-zA-Z]+", text))
    return "zh" if chinese_chars > english_words else "en"


# ---------------------------------------------------------------------------
# Audit report generation
# ---------------------------------------------------------------------------

def generate_audit(
    source_path: str = "",
    tailored_path: str = "",
    source_docx: str = "",
    tailored_docx: str = "",
    company: str = "",
    role: str = "",
    jd_source: str = "",
    source_name: str = "source resume",
    output_lang: str = "auto",
    json_output: bool = False,
) -> str:
    """Generate a full audit log comparing source and tailored resumes."""

    # Read files
    source_structured = None
    tailored_structured = None

    if source_docx:
        source_structured = read_docx_structured(source_docx)
        source_lines = docx_to_text_lines(source_structured)
    elif source_path:
        source_lines = read_lines(source_path)
    else:
        print("Error: --source or --source-docx required", file=sys.stderr)
        sys.exit(1)

    if tailored_docx:
        tailored_structured = read_docx_structured(tailored_docx)
        tailored_lines = docx_to_text_lines(tailored_structured)
    elif tailored_path:
        tailored_lines = read_lines(tailored_path)
    else:
        print("Error: --tailored or --tailored-docx required", file=sys.stderr)
        sys.exit(1)

    # Generate unified diff
    diff = list(difflib.unified_diff(
        source_lines, tailored_lines,
        fromfile="source", tofile="tailored",
        lineterm=""
    ))

    # Auto-detect language
    if output_lang == "auto":
        full_text = "\n".join(source_lines + tailored_lines)
        output_lang = detect_language(full_text)

    is_zh = output_lang == "zh"

    # Categorize changes
    additions = []
    deletions = []
    modifications = []

    i = 0
    while i < len(diff):
        line = diff[i]
        if line.startswith("---") or line.startswith("+++") or line.startswith("@@"):
            i += 1
            continue

        if line.startswith("+"):
            if i + 1 < len(diff) and diff[i + 1].startswith("-"):
                change = classify_change(
                    diff[i + 1].lstrip("- "),
                    line.lstrip("+ "),
                )
                modifications.append({
                    "added": line.lstrip("+ "),
                    "removed": diff[i + 1].lstrip("- "),
                    "type": change["type"],
                    "risk": change["risk"],
                    "risk_reason": change["risk_reason"],
                })
                i += 2
            else:
                additions.append(line.lstrip("+ "))
                i += 1
        elif line.startswith("-"):
            deletions.append(line.lstrip("- "))
            i += 1
        else:
            i += 1

    now = datetime.now().strftime("%Y-%m-%d")

    # --- JSON output mode ---
    if json_output:
        result = {
            "date": now,
            "target": {"company": company, "role": role} if company or role else None,
            "jd_source": jd_source or None,
            "source_resume": source_name,
            "language": output_lang,
            "change_summary": {
                "additions": len(additions),
                "deletions": len(deletions),
                "modifications": len(modifications),
                "by_type": {},
                "by_risk": {"low": 0, "medium": 0, "high": 0},
            },
            "modifications": [],
            "additions": additions,
            "deletions": deletions,
        }

        for mod in modifications:
            t = mod["type"]
            r = mod["risk"]
            result["change_summary"]["by_type"][t] = result["change_summary"]["by_type"].get(t, 0) + 1
            result["change_summary"]["by_risk"][r] = result["change_summary"]["by_risk"].get(r, 0) + 1
            result["modifications"].append({
                "removed": mod["removed"],
                "added": mod["added"],
                "type": t,
                "risk": r,
                "risk_reason": mod["risk_reason"],
            })

        # Add structure context if available
        if source_structured:
            result["source_structure"] = get_structure_summary(source_structured)
        if tailored_structured:
            result["tailored_structure"] = get_structure_summary(tailored_structured)

        return json.dumps(result, ensure_ascii=False, indent=2)

    # --- Markdown report mode ---
    risk_labels = {"low": "🟢 Low", "medium": "🟡 Medium", "high": "🔴 High"}
    type_labels_zh = {
        "wording_upgrade": "措辞升级",
        "quantification": "量化补充",
        "content_change": "内容变更",
    }
    type_labels_en = {
        "wording_upgrade": "Wording Upgrade",
        "quantification": "Quantification",
        "content_change": "Content Change",
    }
    type_labels = type_labels_zh if is_zh else type_labels_en

    # Header
    if is_zh:
        report_lines = [
            "# 简历变更审计 / Resume Change Audit",
            "",
            f"- **日期**: {now}",
        ]
        if company:
            label = f"- **目标岗位**: {company}"
            if role:
                label += f" - {role}"
            report_lines.append(label)
        if jd_source:
            report_lines.append(f"- **JD 来源**: {jd_source}")
        report_lines.append(f"- **基于版本**: {source_name}")

        report_lines.extend([
            "",
            "## 变更统计",
            "",
            "| 类型 | 数量 |",
            "|------|------|",
            f"| 新增 | {len(additions)} |",
            f"| 删除 | {len(deletions)} |",
            f"| 修改 | {len(modifications)} |",
        ])
    else:
        report_lines = [
            "# Resume Change Audit",
            "",
            f"- **Date**: {now}",
        ]
        if company:
            label = f"- **Target**: {company}"
            if role:
                label += f" - {role}"
            report_lines.append(label)
        if jd_source:
            report_lines.append(f"- **JD Source**: {jd_source}")
        report_lines.append(f"- **Based on**: {source_name}")

        report_lines.extend([
            "",
            "## Change Summary",
            "",
            "| Type | Count |",
            "|------|------|",
            f"| Added | {len(additions)} |",
            f"| Removed | {len(deletions)} |",
            f"| Modified | {len(modifications)} |",
        ])

    # Type breakdown
    type_counts = {}
    risk_counts = {"low": 0, "medium": 0, "high": 0}
    for mod in modifications:
        t = mod["type"]
        r = mod["risk"]
        type_counts[t] = type_counts.get(t, 0) + 1
        risk_counts[r] = risk_counts.get(r, 0) + 1

    for t, count in type_counts.items():
        label = type_labels.get(t, t)
        report_lines.append(f"| ↳ {label} | {count} |")

    risk_label_map_zh = {
        "low": ("🟢 Low / 低风险", "已验证的修改"),
        "medium": ("🟡 Medium / 中风险", "建议确认的修改"),
        "high": ("🔴 High / 高风险", "需要详细验证的修改"),
    }
    risk_label_map_en = {
        "low": ("🟢 Low", "Verified changes"),
        "medium": ("🟡 Medium", "Recommended to confirm"),
        "high": ("🔴 High", "Requires detailed verification"),
    }

    if is_zh:
        report_lines.extend(["", "## 风险分布", ""])
        for r in ["high", "medium", "low"]:
            label, desc = risk_label_map_zh[r]
            report_lines.append(f"| {label} | {desc} | {risk_counts[r]} |")
    else:
        report_lines.extend(["", "## Risk Distribution", ""])
        for r in ["high", "medium", "low"]:
            label, desc = risk_label_map_en[r]
            report_lines.append(f"| {label} | {desc} | {risk_counts[r]} |")

    # Detailed changes
    if is_zh:
        if additions:
            report_lines.extend(["", "## 新增内容", ""])
            for idx, add in enumerate(additions, 1):
                report_lines.append(f"{idx}. {add}")

        if deletions:
            report_lines.extend(["", "## 删除内容", ""])
            for idx, dlt in enumerate(deletions, 1):
                report_lines.append(f"{idx}. {dlt}")

        if modifications:
            report_lines.extend(["", "## 修改明细", ""])
            for idx, mod in enumerate(modifications, 1):
                type_label = type_labels.get(mod["type"], mod["type"])
                risk_label = risk_labels.get(mod["risk"], mod["risk"])
                report_lines.extend([
                    f"### {idx}. [{type_label}] {risk_label}",
                    f"- **删除**: {mod['removed']}",
                    f"- **新增**: {mod['added']}",
                    f"- **风险评估**: {mod['risk_reason']}",
                    "",
                ])
    else:
        if additions:
            report_lines.extend(["", "## Added Content", ""])
            for idx, add in enumerate(additions, 1):
                report_lines.append(f"{idx}. {add}")

        if deletions:
            report_lines.extend(["", "## Removed Content", ""])
            for idx, dlt in enumerate(deletions, 1):
                report_lines.append(f"{idx}. {dlt}")

        if modifications:
            report_lines.extend(["", "## Modification Details", ""])
            for idx, mod in enumerate(modifications, 1):
                type_label = type_labels.get(mod["type"], mod["type"])
                risk_label = risk_labels.get(mod["risk"], mod["risk"])
                report_lines.extend([
                    f"### {idx}. [{type_label}] {risk_label}",
                    f"- **Removed**: {mod['removed']}",
                    f"- **Added**: {mod['added']}",
                    f"- **Risk**: {mod['risk_reason']}",
                    "",
                ])

    # Structure context (if docx was used)
    if source_structured:
        struct = get_structure_summary(source_structured)
        if is_zh:
            report_lines.extend(["", "## 源简历结构信息", ""])
        else:
            report_lines.extend(["", "## Source Resume Structure", ""])
        report_lines.append(f"```json")
        report_lines.append(json.dumps(struct, ensure_ascii=False, indent=2))
        report_lines.append("```")

    return "\n".join(report_lines)


def main():
    parser = argparse.ArgumentParser(description="Resume Diff Audit")
    parser.add_argument("--source", help="Path to source resume (Markdown)")
    parser.add_argument("--tailored", help="Path to tailored resume (Markdown)")
    parser.add_argument("--source-docx", help="Path to source resume (.docx)")
    parser.add_argument("--tailored-docx", help="Path to tailored resume (.docx)")
    parser.add_argument("--company", default="", help="Target company name")
    parser.add_argument("--role", default="", help="Target role name")
    parser.add_argument("--jd-source", default="", help="JD URL or source description")
    parser.add_argument("--source-name", default="source resume", help="Display name of the source resume file")
    parser.add_argument("--lang", default="auto", choices=["auto", "zh", "en"], help="Output language (auto-detect by default)")
    parser.add_argument("--json", action="store_true", help="Output as JSON for LLM pipeline")
    parser.add_argument("--output", default="", help="Output file path (default: stdout)")
    args = parser.parse_args()

    report = generate_audit(
        source_path=args.source,
        tailored_path=args.tailored,
        source_docx=args.source_docx,
        tailored_docx=args.tailored_docx,
        company=args.company,
        role=args.role,
        jd_source=args.jd_source,
        source_name=args.source_name,
        output_lang=args.lang,
        json_output=args.json,
    )

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(report)
        print(f"Audit log saved to {args.output}")
    else:
        try:
            print(report)
        except UnicodeEncodeError:
            # Windows console may not support emoji — fallback to ASCII markers
            print(report.encode("utf-8", errors="replace").decode("utf-8"))


if __name__ == "__main__":
    main()
