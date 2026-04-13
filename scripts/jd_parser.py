"""
JD Parser — Feature Extraction Layer (v2)

Role: Script layer of the Hybrid Extraction Engine.
Philosophy: EXTRACT features, don't JUDGE them.
- Scripts grab strings that "look like" degrees, certs, scores, dates
- Scripts do NOT decide if a feature matches — LLM does
- Scripts present raw data from both JD and resume for LLM reasoning

This is intentionally "dumb" — it casts a wide net with generic patterns
rather than trying to enumerate every possible degree/cert/language test.
The LLM layer (SKILL.md) handles all semantic interpretation.

Usage:
    python jd_parser.py "<jd_text>"
    python jd_parser.py --file path/to/jd.txt
    python jd_parser.py --file path/to/jd.txt --json
    python jd_parser.py --file path/to/jd.txt --json --resume path/to/resume.docx
"""

import re
import json
import sys
import argparse
from pathlib import Path

from utils import (
    detect_language,
    check_portfolio_links,
)


# ─── Feature patterns (language/region agnostic) ─────────────
# Design principle: use STRUCTURAL patterns (not word lists)
# E.g., "2-4 uppercase letters = acronym", "number + point + number = score"
# This works across languages, industries, and regions.

# --- Experience (structural: number + year/time unit) ---
EXPERIENCE_RULES = [
    # "at least 3 years", "不少于5年", "minimum 3", "至少3年工作经验"
    (r"(?:at\s+least|minimum|min\.?|no\s+less\s+than|不少于|至少)\s*(\d+)\s*\+?\s*(?:years?|yrs?|年(?:以上)?(?:工作)?(?:经验)?)", "years_min"),
    # "3-5 years", "3至5年"
    (r"(\d+)\s*[-–~到至]\s*(\d+)\s*\+?\s*(?:years?|yrs?|年)", "years_range"),
    # "5 years", "5年经验", "5 years of experience"
    (r"(\d+)\s*\+?\s*(?:years?\s*(?:of\s+)?experience|年(?:以上)?(?:工作)?(?:经验)?)", "years_min"),
    # Seniority level (structural keywords)
    (r"(?:fresh\s*grad|entry\s*level|应届|初级|入门)", "entry_level"),
    (r"(?:senior|staff|principal|高级|资深)", "senior_level"),
    (r"(?:mid[- ]?level|中级)", "mid_level"),
]

# --- Education (structural: common degree-like patterns) ---
# NOT an exhaustive list — just patterns that "look like" degree mentions
DEGREE_RULES = [
    # Doctoral
    (r"(?:Ph\.?\s*D|Doctorate|Doktor|博士|Dr\.\s*-?\s*Ing\.)", "degree_doctoral"),
    # Master-level
    (r"(?:Master|M\.?\s*S\.?|MSc|MA|MPhil|MRes|MBA|MEng|硕士|研究生|Diplôme)", "degree_master"),
    # Bachelor-level
    (r"(?:Bachelor|B\.?\s*S\.?|BSc|BA|BEng|本科|学士|本科)", "degree_bachelor"),
    # Below bachelor
    (r"(?:Associate|Diploma|大专|专科|副学士|Foundation|HND)", "degree_associate"),
]

# --- Language proficiency (structural: test name + optional score) ---
# Feature extraction: grab any "known_test_name + score" pattern
# Known test names: structural pattern (2-6 uppercase letters = test acronym)
# Plus explicit patterns for common tests that don't fit the acronym pattern
LANGUAGE_RULES = [
    # Structured tests with scores: "IELTS 7.5", "TOEFL 105", "JLPT N1"
    (r"\b(IELTS|雅思)\s*[:：]?\s*(\d+(?:\.\d+)?)", "test_score"),
    (r"\b(TOEFL|托福)\s*[:：]?\s*(\d+)", "test_score"),
    (r"\b(TOEIC|托业)\s*[:：]?\s*(\d+)", "test_score"),
    # JLPT: "JLPT N1", "N1 level Japanese", "日本語能力試験 N1"
    (r"\b(?:JLPT|日本語能力試験|日語能力)\s*[:：]?\s*([NnPp][1-5])", "test_score"),
    (r"\b([NnPp][1-5])\s+(?:level|级|日本語|Japanese)", "test_score"),
    # TOPIK: "TOPIK Level 5", "TOPIK 5급", "한국어능력시험"
    (r"\b(TOPIK|한국어능력시험|韩语能力)\s*[:：]?\s*(?:level\s+|级\s*|급\s*)?(\d+)", "test_score"),
    (r"(CET[- ]?(?:4|6))[\s:：]*(?:(\d{2,3})(?:分(?:以上|[++])?)?)?", "test_score"),
    # CEFR levels: "C1", "B2", "A1" — only when near language keywords
    (r"(?:CEFR|CEFR\s*level|欧洲语言共同参考框架)\s*[:：]?\s*([A-C][1-2])", "test_score"),
    (r"(?:level|等级)\s*[:：]?\s*([A-C][1-2])\s*(?:or above|及以上|或以上)?", "test_score"),
    # Proficiency descriptors (no score — just the descriptor)
    (r"(?:fluent|proficient|native|bilingual|business)\s+(?:in\s+)?(?:English|Chinese|Japanese|Korean|French|German|Spanish|Italian|Portuguese|Russian|Arabic|Mandarin|Cantonese|中文|日语|韩语|法语|德语|西班牙语|葡萄牙语|俄语|阿拉伯语)", "language_proficiency"),
    (r"(?:English|英文|英语|Chinese|中文|Mandarin|普通话)\s*(?:as\s+(?:a\s+)?(?:working|first|native)\s+language|working\s+proficiency|流利|熟练|精通|母语)", "language_proficiency"),
]

# --- Certifications (structural: acronym or "certified in X" pattern) ---
# Feature extraction: grab ANY certification-like pattern
CERTIFICATION_RULES = [
    # "Certified in X", "持有X证书", "拥有X认证", "X certified"
    (r"(?:certified?\s+(?:in|by|on|through)|持有.{2,30}(?:证书|认证|资格|执照|牌照)|拥有.{2,30}(?:证书|认证|资格|执照)|possess(?:es)?\s+a\s+\w+\s+certification|licens(?:e|ed|ing))", "certification_generic"),
    # Upper-case acronyms (2-6 letters) near cert-related words
    (r"\b([A-Z]{2,6})\b\s*(?:certified|certification|certificate|资格|证书|认证|执照|license)", "certification_acronym"),
    # Bar/Law: "Bar exam", "admitted to practice"
    (r"(?:Bar\s*(?:exam|admission)?|admitted\s+to\s+practice|律师资格|司法考试|法考)", "certification_legal"),
]

# --- GPA / Academic scores (structural: "GPA/绩点 + number") ---
GPA_RULES = [
    (r"\b(?:GPA|gpa|绩点|CGPA)\s*[:：]?\s*(\d+(?:\.\d+)?)\s*(?:/\s*(\d+(?:\.\d+)?))?", "gpa_score"),
    (r"(?:first\s+class|1st\s+class|upper\s+second|2\s*:\s*1|second\s+class|2\s*:\s*2|third\s+class|一等|二等一|二等二|三等|荣誉学位|Honours?)", "degree_honors"),
]

# --- Work authorization / visa (structural: keyword patterns) ---
WORK_AUTH_RULES = [
    (r"(?:work\s*(?:permit|visa|authorization)|employment\s*(?:pass|visa|authorization)|工作签证|工作许可|就业准证|劳工证|工签)", "work_authorization"),
    (r"(?:citizen|citizenship|permanent\s*resident|PR|公民|永久居民|绿卡|永居)", "citizenship_requirement"),
    (r"(?:Singapore\s*(?:citizen|PR|permanent)|SC\s*\/\s*PR|新加坡公民|新加坡永久|新币)", "citizenship_requirement"),
    (r"(?:must\s+(?:have|hold|possess)\s+(?:valid\s+)?(?:work\s*(?:permit|visa|rights)|active\s+clearance|security\s+clearance))", "work_authorization"),
]

# --- Security clearance (structural) ---
SECURITY_RULES = [
    (r"(?:security\s*clearance|secret\s*clearance|top\s*secret|confidential\s*clearance|安全许可|保密许可|涉密)", "security_clearance"),
]

# --- Location / relocation (structural) ---
LOCATION_RULES = [
    (r"(?:based\s+in|located\s+in|work\s+from|relocate\s+to|onsite|on-site|hybrid|remote(?:\s+work)?|base\s*地|驻场|外派|坐班|出差)", "work_arrangement"),
    (r"(?:willing\s+to\s+relocate|open\s+to\s+relocation|接受调岗|愿意出差|可接受外派)", "relocation_willingness"),
]




# ─── Feature patterns (language/region agnostic) ─────────────


def extract_hard_requirements(text: str) -> dict:
    """
    Extract structural features from JD text.
    Returns list of requirement dicts with type, value, raw_match, and extracted numbers.

    Design: cast a WIDE net — grab anything that structurally looks like
    a requirement. Let LLM decide what's relevant and what matches.
    """
    all_matches = []

    all_rules = (
        EXPERIENCE_RULES + DEGREE_RULES + LANGUAGE_RULES +
        CERTIFICATION_RULES + GPA_RULES + WORK_AUTH_RULES +
        SECURITY_RULES + LOCATION_RULES
    )
    for pattern, label in all_rules:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            all_matches.append({
                "start": match.start(),
                "end": match.end(),
                "value": match.group(0).strip(),
                "label": label,
                "match": match,
            })

    # Sort by position, then by length (longer match first for overlap resolution)
    all_matches.sort(key=lambda m: (m["start"], -(m["end"] - m["start"])))

    # Resolve overlapping matches: keep the first (longest) match in each region
    results = []
    last_end = -1
    seen_keys = set()
    for m in all_matches:
        if m["start"] < last_end:
            continue
        value = m["value"]
        label = m["label"]
        key = f"{label}:{value}"
        if key in seen_keys:
            continue
        seen_keys.add(key)
        last_end = m["end"]
        entry = {
            "type": label,
            "value": value,
            "raw_match": value,
            "position": m["start"],
        }
        entry["jd_number"] = _extract_numbers(value)
        results.append(entry)

    return results


def _extract_numbers(text: str) -> str:
    """Extract the first number found in text. Returns empty string if none."""
    m = re.search(r"(\d+(?:\.\d+)?)", text)
    return m.group(1) if m else ""


def check_resume_match(requirements: list, resume_text: str) -> list:
    """
    Check each extracted feature against resume text.

    Strategy: EXTRACT features from both sides, present raw data.
    Script does NOT decide "match" or "not match" for semantic categories.
    Only structural types (years, test scores, acronyms) get definitive matching.

    Categories:
    - DEFINITIVE matching (script decides): years, test_score with numbers,
      certification_acronym, GPA
    - FEATURE PRESENCE (script reports, LLM decides): degree, language_proficiency,
      certification_generic, work_authorization, citizenship, etc.
    """
    checked = []
    for req in requirements:
        req_type = req["type"]
        jd_number = req.get("jd_number", "")

        # ─── DEFINITIVE: Years of experience ───
        if req_type == "years_min" and jd_number:
            resume_years = _find_year_counts(resume_text)
            if resume_years:
                checked.append({
                    **req,
                    "resume_match": True,
                    "resume_number": str(max(resume_years)),
                    "resume_raw": _find_year_context(resume_text),
                    "note": f"JD requires {jd_number} years; resume mentions up to {max(resume_years)} years (LLM to verify relevance)",
                })
            else:
                # No explicit year count — try date ranges
                date_ranges = _find_date_ranges(resume_text)
                if date_ranges:
                    dr_summary = "; ".join(
                        f"{d['raw']}" + (" (ongoing)" if d["is_ongoing"] else "")
                        for d in date_ranges[:5]
                    )
                    checked.append({
                        **req,
                        "resume_match": False,
                        "resume_number": "",
                        "resume_raw": "",
                        "resume_date_ranges": date_ranges[:5],
                        "note": f"JD requires {jd_number} years; no explicit year count found. "
                                f"Date ranges detected in resume: {dr_summary}. "
                                f"LLM should calculate the span from earliest start to latest end "
                                f"(ongoing entries use current date) to estimate total experience.",
                    })
                else:
                    checked.append({
                        **req,
                        "resume_match": False,
                        "resume_number": "",
                        "resume_raw": "",
                        "note": f"JD requires {jd_number} years; no explicit year count or date ranges found in resume.",
                    })

        # ─── DEFINITIVE: Test scores with numbers (IELTS, TOEFL, TOEIC, etc.) ───
        elif req_type == "test_score":
            # Extract the test name (first word or acronym) and search for it
            test_name = _extract_test_name(req["raw_match"])
            resume_info = _find_test_in_resume(resume_text, test_name, req["raw_match"])
            if resume_info["found"]:
                checked.append({
                    **req,
                    "resume_match": True,
                    "resume_number": resume_info.get("number", ""),
                    "resume_raw": resume_info["raw"],
                    "note": f"JD mentions '{req['value']}'; resume shows '{resume_info['raw']}' (LLM to verify if resume score meets JD requirement)",
                })
            else:
                checked.append({
                    **req,
                    "resume_match": False,
                    "resume_number": "",
                    "resume_raw": "",
                    "note": f"JD mentions '{req['value']}'; not found in resume",
                })

        # ─── DEFINITIVE: Certification acronyms ───
        elif req_type == "certification_acronym":
            # Extract the uppercase acronym
            acro_m = re.search(r"\b([A-Z]{2,6})\b", req["raw_match"])
            cert_name = acro_m.group(1) if acro_m else req["raw_match"]
            found = re.search(re.escape(cert_name), resume_text, re.IGNORECASE) is not None
            if found:
                checked.append({
                    **req,
                    "resume_match": True,
                    "resume_number": "",
                    "resume_raw": cert_name,
                    "note": f"Certification acronym '{cert_name}' found in resume",
                })
            else:
                checked.append({
                    **req,
                    "resume_match": False,
                    "resume_number": "",
                    "resume_raw": "",
                    "note": f"Certification acronym '{cert_name}' not found in resume",
                })

        # ─── DEFINITIVE: GPA ───
        elif req_type == "gpa_score":
            resume_gpa = _find_gpa_in_resume(resume_text)
            if resume_gpa["found"]:
                checked.append({
                    **req,
                    "resume_match": True,
                    "resume_number": resume_gpa.get("gpa", ""),
                    "resume_raw": resume_gpa["raw"],
                    "note": f"JD mentions GPA; resume shows '{resume_gpa['raw']}' (LLM to verify scale compatibility: JD may use 4.0, resume may use 5.0 or percentage)",
                })
            else:
                checked.append({
                    **req,
                    "resume_match": False,
                    "resume_number": "",
                    "resume_raw": "",
                    "note": "JD mentions GPA; no GPA found in resume",
                })

        # ─── FEATURE PRESENCE: Degrees ───
        elif req_type.startswith("degree_"):
            degree_level = req_type  # e.g., "degree_doctoral", "degree_master"
            found = _find_degree_in_resume(resume_text, req["raw_match"])
            if found["found"]:
                checked.append({
                    **req,
                    "resume_match": True,
                    "resume_number": "",
                    "resume_raw": found["raw"],
                    "note": f"Degree-related text '{found['raw']}' found in resume (LLM to verify it's the same degree level as JD requires)",
                })
            else:
                checked.append({
                    **req,
                    "resume_match": False,
                    "resume_number": "",
                    "resume_raw": "",
                    "note": f"JD mentions '{req['value']}'; no matching degree found in resume",
                })

        # ─── FEATURE PRESENCE: Everything else ───
        else:
            # Generic: extract keywords and check presence
            keywords = _extract_keywords(req["raw_match"])
            found = any(
                re.search(re.escape(kw), resume_text, re.IGNORECASE)
                for kw in keywords if len(kw) >= 2
            )
            if found:
                checked.append({
                    **req,
                    "resume_match": True,
                    "resume_number": "",
                    "resume_raw": ", ".join(keywords),
                    "note": f"Keyword(s) '{', '.join(keywords)}' found in resume (LLM to verify semantic match)",
                })
            else:
                checked.append({
                    **req,
                    "resume_match": False,
                    "resume_number": "",
                    "resume_raw": "",
                    "note": f"Keyword(s) '{', '.join(keywords)}' not found in resume",
                })

    return checked


# ─── Feature extraction helpers (resume side) ─────────────

def _find_year_counts(text: str) -> list:
    """Extract all year counts mentioned in resume text."""
    patterns = [
        r"(\d+)\s*\+?\s*(?:years?|yrs?|年(?:以上)?(?:工作)?(?:经验)?)",
        r"(?:over|more\s+than|超过|多于)\s*(\d+)\s*(?:years?|年)",
    ]
    numbers = []
    for p in patterns:
        for m in re.finditer(p, text, re.IGNORECASE):
            num = int(m.group(1))
            if num < 50:  # Sanity check: ignore years like "1990"
                numbers.append(num)
    return numbers


def _find_date_ranges(text: str) -> list:
    """Extract date ranges from resume text (e.g., '2021.01 - 2024.03', 'Jan 2020 - present').
    Returns list of {"start": str, "end": str, "raw": str, "is_ongoing": bool}.
    """
    ranges = []
    # Numeric: 2021.01 - 2024.03, 2021/01 - 2024/03, 2021-01 - 2024-03
    p1 = re.finditer(
        r"(\d{4})[.\-/年](\d{1,2})[月]?\s*[-–—~至到]\s*(\d{4})[.\-/年](\d{1,2})[月]?",
        text
    )
    for m in p1:
        ranges.append({
            "start": f"{m.group(1)}.{m.group(2).zfill(2)}",
            "end": f"{m.group(3)}.{m.group(4).zfill(2)}",
            "raw": m.group(0).strip(),
            "is_ongoing": False,
        })
    # Ongoing: 2021.01 - present/till now/至今
    p2 = re.finditer(
        r"(\d{4})[.\-/年](\d{1,2})[月]?\s*[-–—~至到]\s*(?:present|till\s+now|current|至今|现在|至今)",
        text, re.IGNORECASE
    )
    for m in p2:
        ranges.append({
            "start": f"{m.group(1)}.{m.group(2).zfill(2)}",
            "end": "",
            "raw": m.group(0).strip(),
            "is_ongoing": True,
        })
    # Month name: Jan 2020 - Dec 2024
    p3 = re.finditer(
        r"([A-Za-z]{3,9})\s+(\d{4})\s*[-–—~至到]\s*([A-Za-z]{3,9})\s+(\d{4})",
        text, re.IGNORECASE
    )
    for m in p3:
        ranges.append({
            "start": f"{m.group(2)}.{m.group(1)}",
            "end": f"{m.group(4)}.{m.group(3)}",
            "raw": m.group(0).strip(),
            "is_ongoing": False,
        })
    # Month name ongoing: Jan 2020 - present
    p4 = re.finditer(
        r"([A-Za-z]{3,9})\s+(\d{4})\s*[-–—~至到]\s*(?:present|till\s+now|current|至今|现在)",
        text, re.IGNORECASE
    )
    for m in p4:
        ranges.append({
            "start": f"{m.group(2)}.{m.group(1)}",
            "end": "",
            "raw": m.group(0).strip(),
            "is_ongoing": True,
        })
    return ranges


def _find_year_context(text: str) -> str:
    """Return surrounding context of the first year count found."""
    m = re.search(
        r"(.{0,40}?(\d+)\s*\+?\s*(?:years?|yrs?|年(?:以上)?(?:工作)?(?:经验)?).{0,20}?)",
        text, re.IGNORECASE
    )
    return m.group(1).strip() if m else ""


def _extract_test_name(raw_match: str) -> str:
    """Extract the test name from a raw match. E.g., 'IELTS 7.0' → 'IELTS'."""
    m = re.match(r"([A-Za-z\u4e00-\u9fff]+)", raw_match)
    return m.group(1).strip() if m else raw_match.strip()


def _find_test_in_resume(resume_text: str, test_name: str, jd_raw: str) -> dict:
    """Find any mention of a test in resume text, including score if available."""
    # Search by test name (case-insensitive)
    pattern = rf"{re.escape(test_name)}[\s:：\-–]*(\d+(?:\.\d+)?)?"
    m = re.search(pattern, resume_text, re.IGNORECASE)
    if m:
        return {
            "found": True,
            "number": m.group(1) if m.group(1) else "",
            "raw": m.group(0).strip(),
        }

    # For CET: also check Chinese form
    if "CET" in test_name.upper():
        cet_m = re.search(r"CET[- ]?(4|6)", jd_raw, re.IGNORECASE)
        cet_level = cet_m.group(1) if cet_m else ""
        if cet_level:
            zh_level = "四" if cet_level == "4" else "六"
            zh_pattern = rf"英语{zh_level}(?:级)?[\s:：]*(?:(\d{{2,3}})(?:分|[+])?)?"
            m2 = re.search(zh_pattern, resume_text)
            if m2:
                return {
                    "found": True,
                    "number": m2.group(1) if m2.group(1) else "",
                    "raw": m2.group(0).strip(),
                }

    return {"found": False, "number": "", "raw": ""}


def _find_gpa_in_resume(text: str) -> dict:
    """Find GPA in resume text, including scale if available."""
    # "GPA 3.8/4.0", "GPA: 3.8", "绩点 4.5/5.0"
    m = re.search(
        r"(?:GPA|gpa|CGPA|cgpa|绩点)\s*[:：]?\s*(\d+(?:\.\d+)?)\s*(?:/\s*(\d+(?:\.\d+)?))?",
        text
    )
    if m:
        return {
            "found": True,
            "gpa": m.group(1),
            "raw": m.group(0).strip(),
        }
    return {"found": False, "gpa": "", "raw": ""}


def _find_degree_in_resume(text: str, jd_raw: str) -> dict:
    """
    Find degree-related text in resume.
    Searches for any substring of the JD match that's 3+ chars long.
    This handles unknown degree types (e.g., "Diplôme d'Ingénieur", "MPhil").
    """
    # Try matching progressively shorter substrings (longest first)
    # This handles cases where JD says "Master's degree" but resume says "MSc"
    words = re.findall(r"[A-Za-z\u4e00-\u9fff']{3,}", jd_raw)
    for word in sorted(words, key=len, reverse=True):
        if re.search(re.escape(word), text, re.IGNORECASE):
            context_m = re.search(
                rf"(.{{0,30}}?{re.escape(word)}.{{0,20}}?)",
                text, re.IGNORECASE
            )
            return {
                "found": True,
                "raw": context_m.group(1).strip() if context_m else word,
            }
    return {"found": False, "raw": ""}


def _extract_keywords(text: str) -> list:
    """Extract meaningful keywords from text for matching."""
    words = re.findall(r"[a-zA-Z\u4e00-\u9fff]+", text)
    return [w for w in words if len(w) >= 2]


# ─── Main parsing ─────────────

def parse_jd(text: str, resume_text: str = None) -> dict:
    """
    Main parsing function. Extracts structural features from JD
    and optionally checks against resume text.
    """
    lang = detect_language(text)
    requirements = extract_hard_requirements(text)

    # Sort by position in JD (top of JD = more important)
    requirements.sort(key=lambda x: x["position"])

    # Remove internal fields
    for req in requirements:
        req.pop("position", None)
        if not resume_text:
            req.pop("jd_number", None)

    # Check against resume if provided
    if resume_text:
        requirements = check_resume_match(requirements, resume_text)

    # Build summary
    pdf_reader_used = _pdf_reader_used
    summary = {
        "total_requirements": len(requirements),
        "categories": list(set(r["type"] for r in requirements)),
    }

    if resume_text:
        matched = [r for r in requirements if r.get("resume_match")]
        unmatched = [r for r in requirements if not r.get("resume_match")]
        summary["matched"] = len(matched)
        summary["unmatched"] = len(unmatched)

        # hard_gaps: raw data for LLM — script does NOT judge severity
        summary["hard_gaps"] = []
        for r in unmatched:
            gap = {
                "type": r["type"],
                "jd_value": r["value"],
            }
            if r.get("jd_number"):
                gap["jd_number"] = r["jd_number"]
            if r.get("resume_number"):
                gap["resume_number"] = r["resume_number"]
            summary["hard_gaps"].append(gap)

        # Portfolio / links check (script layer)
        portfolio = check_portfolio_links(text, resume_text)
        summary["portfolio"] = portfolio

    output = {
        "language": lang,
        "hard_requirements": requirements,
        "summary": summary,
    }

    if pdf_reader_used:
        output["pdf_reader_used"] = pdf_reader_used

    return output


# ─── File reading ─────────────

def read_resume_text(path: str) -> str:
    """
    Read resume file and return plain text.

    Format support (in priority order):
    - .txt / .md: direct text read
    - .docx: python-docx (paragraphs + tables)
    - .pdf: pdfplumber > PyPDF2 > pdftotext CLI > raw fallback
    """
    path = Path(path)
    if not path.exists():
        print(json.dumps({"error": f"Resume file not found: {path}"}, ensure_ascii=False))
        sys.exit(1)

    suffix = path.suffix.lower()

    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8")

    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except ImportError:
            return _read_fallback(path)

    if suffix == ".pdf":
        return _read_pdf(path)

    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return ""


# Track which PDF reader was used (module-level for reporting)
_pdf_reader_used = None


def _read_pdf(path: Path) -> str:
    """Read PDF with layered fallback: pdfplumber > PyPDF2 > pdftotext CLI > raw."""
    global _pdf_reader_used

    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
        if texts:
            _pdf_reader_used = "pdfplumber"
            return "\n".join(texts)
    except ImportError:
        pass
    except Exception:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        if texts:
            _pdf_reader_used = "PyPDF2"
            return "\n".join(texts)
    except ImportError:
        pass
    except Exception:
        pass

    try:
        import subprocess
        result = subprocess.run(
            ["pdftotext", str(path), "-"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            _pdf_reader_used = "pdftotext"
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    _pdf_reader_used = "raw_fallback"
    return _read_fallback(path)


def _read_fallback(path: Path) -> str:
    """Fallback: read binary file as text, skip errors."""
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


# ─── Output ─────────────

def print_report(analysis: dict):
    """Print a human-readable report."""
    lang = analysis.get("language", "en")
    is_zh = lang == "zh"

    sep = "=" * 60
    title = "JD 硬性要求分析" if is_zh else "JD Hard Requirements Analysis"
    print(sep)
    print(title)
    print(sep)

    reqs = analysis.get("hard_requirements", [])
    if not reqs:
        no_req = "未检测到硬性要求" if is_zh else "No hard requirements detected"
        print(f"\n  {no_req}")
    else:
        for req in reqs:
            match_icon = "✅" if req.get("resume_match") else "❌" if req.get("resume_match") is False else "—"
            line = f"  {match_icon} [{req['type']}] {req['value']}"
            if req.get("note"):
                line += f"\n       └─ {req['note']}"
            print(line)

    summary = analysis.get("summary", {})
    total_label = "要求总数" if is_zh else "Total requirements"
    print(f"\n## {total_label}: {summary.get('total_requirements', len(reqs))}")

    if "matched" in summary:
        matched_label = "已匹配" if is_zh else "Matched"
        unmatched_label = "未匹配" if is_zh else "Unmatched"
        print(f"  {matched_label}: {summary['matched']}")
        print(f"  {unmatched_label}: {summary['unmatched']}")

    if summary.get("hard_gaps"):
        gap_label = "⚠️ 硬性差距（待 LLM 判断严重程度）" if is_zh else "⚠️ Hard Gaps (for LLM assessment)"
        print(f"\n## {gap_label}")
        for gap in summary["hard_gaps"]:
            gap_line = f"  • [{gap['type']}] JD: {gap['jd_value']}"
            if gap.get("jd_number"):
                gap_line += f" (number: {gap['jd_number']})"
            if gap.get("resume_number"):
                gap_line += f" | Resume: {gap['resume_number']}"
            print(gap_line)

    print(f"\n{sep}")


def validate_output(analysis: dict) -> bool:
    """Validate output matches expected JSON Schema."""
    required_fields = ["language", "hard_requirements", "summary"]
    for field in required_fields:
        if field not in analysis:
            return False
    if not isinstance(analysis["hard_requirements"], list):
        return False
    for req in analysis["hard_requirements"]:
        if not all(k in req for k in ["type", "value"]):
            return False
    return True


def main():
    parser = argparse.ArgumentParser(description="JD Feature Extractor (Hybrid Engine - Script Layer)")
    parser.add_argument("input", help="JD text or file path")
    parser.add_argument("--file", action="store_true", help="Read JD from file")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    parser.add_argument("--resume", type=str, default=None, help="Resume file path for match checking")
    args = parser.parse_args()

    if args.file:
        with open(args.input, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        text = args.input

    resume_text = None
    if args.resume:
        resume_text = read_resume_text(args.resume)

    analysis = parse_jd(text, resume_text)

    if not validate_output(analysis):
        error = {"error": "Output validation failed", "raw_output": str(analysis)}
        print(json.dumps(error, ensure_ascii=False, indent=2))
        sys.exit(1)

    if args.json:
        print(json.dumps(analysis, ensure_ascii=False, indent=2))
    else:
        print_report(analysis)


if __name__ == "__main__":
    main()
